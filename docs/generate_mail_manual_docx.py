# -*- coding: utf-8 -*-
"""ヴィレッジ メール設定マニュアル（iPhone / Android / パソコン）の DOCX 生成

方針：業務文書として簡潔に。平易な語で書くが、断定的な所要時間や口語表現は用いない。
  - 設定値は冒頭の①に集約し、各手順では繰り返さない
  - 1手順＝1行。専門用語は初出のみカッコで補足
  - パスワードは本文に書かない（別紙「メール設定一覧」で管理）
出力：メール設定マニュアル_iPhone_Android_PC.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0F, 0x35, 0x6E)
GRAY = RGBColor(0x66, 0x66, 0x66)
FONT = "Meiryo UI"

MAIL = "info@village2024.jp"
SERVER = "v2008.coreserver.jp"


def set_font(run, size=10.5, bold=False, color=None, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font)


def para(doc, text="", size=10.5, bold=False, color=None, align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        set_font(p.add_run(text), size, bold, color)
    return p


def heading(doc, text, size=14):
    p = para(doc, text, size=size, bold=True, color=NAVY, space_after=6)
    p.paragraph_format.space_before = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '2'); bottom.set(qn('w:color'), 'F5A623')
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)


def make_table(doc, headers, rows, widths=None, head_fill='0F356E'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].clear()
        set_font(hdr[i].paragraphs[0].add_run(h), 10, True, RGBColor(0xFF, 0xFF, 0xFF))
        shade(hdr[i], head_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].clear()
            set_font(cells[i].paragraphs[0].add_run(str(val)), 9.5, False)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def note(doc, text):
    return para(doc, text, size=9, color=GRAY, space_after=2)


def steps(doc, items):
    for i, s in enumerate(items, 1):
        para(doc, f"{i}. {s}", size=10.5, space_after=2)


def build():
    doc = Document()
    st = doc.styles['Normal']; st.font.name = FONT; st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    para(doc, "メール設定マニュアル", size=18, bold=True, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "iPhone ／ Android ／ パソコン", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "株式会社ヴィレッジ　御中", size=11, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    para(doc, f"対象メール：@village2024.jp（CoreServer） ／ 作成 2026-07-07・改訂 2026-07-31",
         size=9.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    para(doc, "本書は、会社のメールを各端末で送受信するための設定手順書です。", size=10.5, space_after=2)
    para(doc, "①の設定情報をご用意のうえ、②〜④のうち該当する端末の手順をご覧ください。",
         size=10.5, space_after=2)

    # ---------------------------------------------------------------- ①
    heading(doc, "① 設定情報（最初にご確認ください）")
    make_table(doc, ["項目", "入力する値"], [
        ["メールアドレス", MAIL],
        ["パスワード", "メールのパスワード（別紙「メール設定一覧」に記載）"],
        ["サーバー名", f"{SERVER}（受信・送信とも同じ）"],
        ["ユーザー名", f"{MAIL}（メールアドレスをそのまま入力）"],
        ["メールの種類", "IMAP（アイマップ）"],
    ], widths=[3.6, 11.6])
    note(doc, f"※サーバー名は {SERVER} です。独自ドメイン village2024.jp ではありません。"
              "誤入力の多い項目のため、ご注意ください。")
    note(doc, "※IMAP は、サーバー上でメールを管理する方式です。複数の端末で同じ受信トレイを共有でき、"
              "機種変更後も過去のメールを参照できます。")

    para(doc, "")
    para(doc, "ポート番号・暗号化（CoreServer 標準値・変更不要）", size=11, bold=True,
         color=NAVY, space_after=3)
    make_table(doc, ["", "ポート番号", "暗号化", "認証（ログイン）"], [
        ["受信", "993", "SSL/TLS＝オン", "あり"],
        ["送信", "465", "SSL/TLS＝オン", "あり（受信と同じID・パスワード）"],
    ], widths=[2.2, 3.4, 4.4, 5.2])
    note(doc, "※接続できない場合の代替値：受信 143／送信 587（暗号化は STARTTLS を選択）。"
              "なお 143 と SSL/TLS の組み合わせでは接続できません。")

    # ---------------------------------------------------------------- ②
    heading(doc, "② iPhone・iPad の設定")
    para(doc, "iPhone に標準搭載の「メール」アプリを使用します。", size=10.5, space_after=3)
    steps(doc, [
        "「設定」（歯車のアイコン）を開きます。",
        "「アプリ」→「メール」を選択します。（iOS のバージョンにより「メール」が直下にある場合があります）",
        "「メールアカウント」→「アカウントを追加」を選択します。",
        "「その他」→「メールアカウントを追加」を選択します。",
        "名前（例：ヴィレッジ）・メール・パスワード・説明（例：会社メール）を入力し「次へ」。",
        "画面上部のタブが「IMAP」になっていることを確認します。",
        "「受信メールサーバ」「送信メールサーバ」の両方に、①のサーバー名・ユーザ名・パスワードを入力します。",
        "右上の「次へ」→「保存」を選択します。確認処理に時間を要する場合があります。",
    ])
    note(doc, "※送信側が「オプション」と表示されていても、ユーザ名とパスワードは必ず入力してください。"
              "未入力の場合、受信はできても送信ができません。")
    note(doc, "※「サーバの識別情報を検証できません」と表示された場合は「続ける」を選択してください。")

    # ---------------------------------------------------------------- ③
    heading(doc, "③ Android の設定")
    para(doc, "機種による差異が少ない「Gmail」アプリを使用します。"
              "（Gmail のアドレスをお持ちでない場合も設定できます）", size=10.5, space_after=3)
    steps(doc, [
        "「Gmail」アプリを開きます。",
        "右上のアカウントアイコン →「別のアカウントを追加」を選択します。",
        "「その他」を選択し、メールアドレスを入力して「次へ」。",
        "「個人用（IMAP）」を選択し、パスワードを入力して「次へ」。",
        "「受信サーバーの設定」に①のユーザー名・パスワード・サーバー名を入力し、ポート993／SSL を選択して「次へ」。",
        "「送信サーバーの設定」で「ログインが必要」をオンにし、同じ値を入力、ポート465／SSL を選択して「次へ」。",
        "同期頻度などの確認画面は、そのまま「次へ」。",
        "名前（例：ヴィレッジ）を入力し「次へ」を選択します。",
    ])
    note(doc, "※機種により「個人用（IMAP）」が「IMAP」と表示される場合があります。同一の設定です。")

    # ---------------------------------------------------------------- ④
    heading(doc, "④ パソコン（Windows）の設定")
    para(doc, "無料のメールソフト「Thunderbird（サンダーバード）」を使用する手順です。"
              "Outlook をご利用の場合は、下段の「Outlook の場合」をご覧ください。", size=10.5, space_after=3)

    para(doc, "Thunderbird のインストール（未導入の場合）", size=10.5, bold=True, color=NAVY, space_after=2)
    steps(doc, [
        "Edge などのブラウザで「Thunderbird ダウンロード」と検索します。",
        "公式サイト（thunderbird.net）から「無料ダウンロード」を選択します。",
        "ダウンロードしたファイルを開き、画面の案内に従って進めます。",
    ])

    para(doc, "アカウントの追加", size=10.5, bold=True, color=NAVY, space_after=2)
    steps(doc, [
        "Thunderbird を起動します。",
        "「新しいアカウントを作成」→「メール」を選択します。",
        "名前（例：ヴィレッジ）・メールアドレス・パスワードを入力し「続ける」。",
        "自動設定では正しく検出されないため「手動設定」を開き、下表のとおり入力します。",
        "「再テスト」→「完了」を選択します。受信トレイが表示されれば設定完了です。",
    ])
    make_table(doc, ["項目", "受信サーバー", "送信サーバー"], [
        ["方式", "IMAP", "SMTP"],
        ["サーバー名", SERVER, f"{SERVER}（受信と同じ）"],
        ["ポート", "993", "465"],
        ["接続の保護", "SSL/TLS", "SSL/TLS"],
        ["認証方式", "通常のパスワード認証", "通常のパスワード認証"],
        ["ユーザー名", MAIL, MAIL],
    ], widths=[3.2, 6.0, 6.0])
    note(doc, "※2台目以降のパソコンも同じ手順です。IMAP のため、受信トレイは全端末で共有されます。")

    para(doc, "Outlook の場合", size=10.5, bold=True, color=NAVY, space_after=2)
    steps(doc, [
        "「ファイル」→「アカウントの追加」を選択します。",
        "メールアドレスを入力し、「詳細オプション」→「自分で自分のアカウントを手動で設定」にチェックして「接続」。",
        "種類は「IMAP」を選択します。",
        f"受信・送信ともサーバー {SERVER}／受信993・送信465／暗号化 SSL/TLS を入力します。",
        "パスワードを入力し「接続」を選択します。受信トレイが表示されれば設定完了です。",
    ])

    # ---------------------------------------------------------------- ⑤
    heading(doc, "⑤ 送受信テスト")
    para(doc, "設定後、受信・送信の両方をご確認ください。", size=10.5, space_after=3)
    para(doc, "1. 受信：他のメールアドレス（個人の Gmail など）から会社アドレス宛に送信し、受信を確認します。",
         size=10.5, space_after=2)
    para(doc, "2. 送信：会社アドレスから他のメールアドレス宛に送信し、到着を確認します。",
         size=10.5, space_after=2)
    note(doc, "※両方が確認できれば設定完了です。いずれかに失敗する場合は⑥をご確認ください。")

    # ---------------------------------------------------------------- ⑥
    heading(doc, "⑥ トラブル対応")
    make_table(doc, ["症状・ご質問", "確認する項目"], [
        ["受信はできるが、送信できない", "送信の「認証（ログインが必要）」がオフになっていないか確認。または送信ポートを 465 ⇔ 587 で切替。"],
        ["送信はできるが、受信できない", "受信のサーバー名・ポート993・暗号化（オン）・パスワードを再確認。"],
        ["「接続できません」と表示される", f"サーバー名の誤入力が主な原因です。{SERVER} で相違ないか確認。"],
        ["ポート143 で接続できない", "143 の場合は暗号化を STARTTLS に変更。推奨は 993＋SSL/TLS。"],
        ["パスワードが不明", "CoreServer の管理画面で再設定できます。"],
        ["複数端末で同じメールを参照したい", "本書のとおり IMAP で設定すれば、受信トレイは共有されます。"],
        ["機種変更でメールは消えるか", "消えません。新しい端末で同じ設定を行えば、過去のメールも参照できます。"],
    ], widths=[5.4, 9.8])

    para(doc, "")
    para(doc, "作成：株式会社ヴィレッジ コーポレートサイト制作チーム（ダイマグ）", size=9,
         color=GRAY, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    para(doc, "CoreServer の標準設定に基づいています。仕様変更で項目名が変わる場合があります。", size=9,
         color=GRAY, align=WD_ALIGN_PARAGRAPH.RIGHT)

    out = __file__.replace("generate_mail_manual_docx.py",
                           "メール設定マニュアル_iPhone_Android_PC.docx")
    doc.save(out)
    print("saved:", out)


if __name__ == "__main__":
    build()
